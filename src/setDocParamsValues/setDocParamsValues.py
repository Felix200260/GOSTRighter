import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def generate_unique_filename(path):
    base, ext = os.path.splitext(path)
    count = 1
    new_path = f"{base}_modified_{count}{ext}"
    while os.path.exists(new_path):
        count += 1
        new_path = f"{base}_modified_{count}{ext}"
    return new_path

def apply_new_document_params(original_doc, font_size, indent_size, line_spacing, font_name, margin_left, margin_right, margin_top, margin_bottom):
    new_doc = Document()
    
    # Установка стилей и параметров для нового документа
    style = new_doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.italic = False

    for section in new_doc.sections:
        section.left_margin = Mm(margin_left)
        section.right_margin = Mm(margin_right)
        section.top_margin = Mm(margin_top)
        section.bottom_margin = Mm(margin_bottom)

    def copy_paragraph(paragraph, new_doc):
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.style = style
        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        new_paragraph.paragraph_format.left_indent = Cm(0)
        new_paragraph.paragraph_format.right_indent = Cm(0)
        new_paragraph.paragraph_format.first_line_indent = Cm(indent_size)
        new_paragraph.paragraph_format.space_before = Pt(0)
        new_paragraph.paragraph_format.space_after = Pt(0)
        new_paragraph.paragraph_format.line_spacing = line_spacing

        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.font.name = font_name
            new_run.font.size = Pt(font_size)
            new_run.font.italic = False
            new_run.font.bold = run.bold
            new_run.font.underline = run.underline
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.highlight_color:
                new_run.font.highlight_color = run.font.highlight_color
            if run.font.strike:
                new_run.font.strike = run.font.strike

    def copy_picture(paragraph, new_doc):
        for run in paragraph.runs:
            if run.element.xml.find('w:drawing') != -1:
                new_paragraph = new_doc.add_paragraph()
                new_paragraph.style = style
                new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                new_paragraph.paragraph_format.left_indent = Cm(0)
                new_paragraph.paragraph_format.right_indent = Cm(0)
                new_paragraph.paragraph_format.first_line_indent = Cm(indent_size)
                new_paragraph.paragraph_format.space_before = Pt(0)
                new_paragraph.paragraph_format.space_after = Pt(0)
                new_paragraph.paragraph_format.line_spacing = line_spacing
                new_paragraph._element.append(run.element)

    # Копирование содержимого из оригинального документа в новый
    for paragraph in original_doc.paragraphs:
        copy_paragraph(paragraph, new_doc)
        copy_picture(paragraph, new_doc)

    for table in original_doc.tables:
        new_table = new_doc.add_table(rows=0, cols=len(table.columns))
        for row in table.rows:
            new_row = new_table.add_row().cells
            for idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    copy_paragraph(paragraph, new_row[idx])

    return new_doc

def apply_document_params(doc_path, params):
    original_doc = Document(doc_path)

    font_size = params.get('font_size', 12)
    indent_size = params.get('indent_size', 1.25)
    line_spacing = params.get('line_spacing', 1.5)
    font_names = params.get('recommended_fonts', ['Times New Roman'])
    font_name = font_names[0] if font_names else 'Times New Roman'

    margin_left = params.get('margin_size_left', 3.0)
    margin_right = params.get('margin_size_right', 1.5)
    margin_top = params.get('margin_size_top', 2.0)
    margin_bottom = params.get('margin_size_bottom', 2.0)

    print(f"Получены такие параметры документа: шрифт {font_name}, размер шрифта {font_size}, отступ {indent_size}, межстрочный интервал {line_spacing}")
    print(f"Размер поля left: {margin_left}")
    print(f"Размер поля right: {margin_right}")
    print(f"Размер поля top: {margin_top}")
    print(f"Размер поля bottom: {margin_bottom}")

    new_doc = apply_new_document_params(
        original_doc,
        font_size=font_size,
        indent_size=indent_size,
        line_spacing=line_spacing,
        font_name=font_name,
        margin_left=margin_left,
        margin_right=margin_right,
        margin_top=margin_top,
        margin_bottom=margin_bottom
    )

    modified_path = generate_unique_filename(doc_path)
    new_doc.save(modified_path)
    print(f"Документ сохранен как {modified_path}")

    return modified_path

if __name__ == "__main__":
    doc_path = r'C:/Users/felix/YandexDisk-korchevskyfelix/Programming/Programming/Python/GOSTRighter/tests/word/testFileWord.docx'
    
    params = {
        'font_size': 12,
        'indent_size': 1.25,
        'line_spacing': 1.5,
        'recommended_fonts': ['Times New Roman'],
        'margin_size_left': 30,
        'margin_size_right': 15,
        'margin_size_top': 20,
        'margin_size_bottom': 20
    }

    modified_doc_path = apply_document_params(doc_path, params)
